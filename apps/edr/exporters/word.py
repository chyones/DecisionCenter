"""Word (.docx) exporter — for formal executive reports."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from apps.edr.graph.report_policy import (
    SEC_CONTRACTUAL,
    SEC_DELAY_ANALYSIS,
    SEC_FINANCIAL_SNAPSHOT,
    SEC_ROOT_CAUSES,
    policy_for,
)


_BRAND_BLUE = RGBColor(0x1F, 0x38, 0x64)


def to_word(report: dict) -> bytes:
    """Generate a formal Word document from the canonical JSON report dict."""
    doc = Document()

    request_id = report.get("request_id", "N/A")
    project_code = report.get("project_code") or "N/A"
    pid = report.get("project_identity") or {}
    project_name = pid.get("project_name") or project_code
    report_type = report.get("report_type", "executive_decision")
    report_title = report_type.replace("_", " ").title()
    query = report.get("query") or report.get("question", "N/A")
    language = report.get("language", "en")
    qg_status = report.get("quality_gate_status", "not_run")
    generated_at = datetime.now(timezone.utc).isoformat(timespec="seconds")
    policy = policy_for(report_type)
    section_no = 0

    def _heading(title: str) -> None:
        nonlocal section_no
        section_no += 1
        doc.add_heading(f"{section_no}. {title}", level=1)

    # Cover heading
    title_para = doc.add_heading(f"{report_title} — {project_name} — {project_code}", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    meta_table = doc.add_table(rows=7, cols=2)
    meta_table.style = "Table Grid"
    for i, (key, val) in enumerate(
        [
            ("Request ID", request_id),
            ("Project Name", project_name),
            ("Project Code", project_code),
            ("Report Type", report_type),
            ("Language", language),
            ("Quality Gate", qg_status),
            ("Generated", generated_at),
        ]
    ):
        meta_table.rows[i].cells[0].text = key
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        meta_table.rows[i].cells[1].text = val
    doc.add_paragraph()
    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph()

    # 1. Executive Summary
    _heading("Executive Summary")
    summary = report.get("executive_summary", [])
    if isinstance(summary, str):
        doc.add_paragraph(summary)
    elif summary:
        for item in summary:
            if isinstance(item, dict):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item.get("claim", ""))
                p.add_run(f"  [Confidence: {item.get('confidence', 'medium')}]").italic = True
                refs = ", ".join(item.get("evidence_ids", []))
                if refs:
                    p.add_run(f" — {refs}").italic = True
    else:
        doc.add_paragraph("No summary available.")

    # Financial Snapshot — Odoo
    if policy.renders(SEC_FINANCIAL_SNAPSHOT):
        _heading("Financial Snapshot — Odoo")
        fs = report.get("financial_snapshot") or {}
        if isinstance(fs, dict):
            budget = fs.get("budget") or {}
            actual = fs.get("actual_cost") or {}
            variance = fs.get("variance") or {}
            currency = (budget.get("currency") if isinstance(budget, dict) else None) or "AED"

            fin_table = doc.add_table(rows=4, cols=3)
            fin_table.style = "Table Grid"
            for i, h in enumerate(["Item", "Value", "Source"]):
                fin_table.rows[0].cells[i].text = h
                fin_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            def _fill_fin_row(row, label: str, node: dict) -> None:
                row.cells[0].text = label
                if isinstance(node, dict):
                    v = node.get("value")
                    c = node.get("currency", currency)
                    row.cells[1].text = f"{v:,.2f} {c}" if v is not None else "Not available"
                    row.cells[2].text = node.get("evidence_id") or "—"
                else:
                    row.cells[1].text = "Not available"
                    row.cells[2].text = "—"

            _fill_fin_row(fin_table.rows[1], "Budget", budget)
            _fill_fin_row(fin_table.rows[2], "Actual Cost", actual)

            fin_table.rows[3].cells[0].text = "Variance"
            if isinstance(variance, dict):
                v = variance.get("value")
                c = variance.get("currency", currency)
                formula = variance.get("formula", "")
                val_str = f"{v:,.2f} {c}" if v is not None else "Not available"
                fin_table.rows[3].cells[1].text = (
                    f"{val_str}{' (' + formula + ')' if formula else ''}"
                )
            else:
                fin_table.rows[3].cells[1].text = "Not available"
            fin_table.rows[3].cells[2].text = "—"
        else:
            doc.add_paragraph("Financial data not available.")
        doc.add_paragraph()

    # 3–7. Findings sections
    section_specs = [
        ("Key Findings", "key_findings"),
        ("Recommended Actions — Proposal Only", "recommended_actions"),
    ]
    if policy.renders(SEC_ROOT_CAUSES):
        section_specs[1:1] = [
            ("Root Causes", "root_causes"),
        ]
    if policy.renders(SEC_DELAY_ANALYSIS):
        insert_at = len(section_specs) - 1
        section_specs[insert_at:insert_at] = [
            ("Delay Analysis", "delay_analysis"),
        ]
    if policy.renders(SEC_CONTRACTUAL):
        insert_at = len(section_specs) - 1
        section_specs[insert_at:insert_at] = [
            ("Contractual / Commercial Implications", "contractual_implications"),
        ]
    for heading, key in section_specs:
        _heading(heading)
        items = report.get(key, [])
        if items:
            for item in items:
                if isinstance(item, dict):
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(item.get("text", str(item)))
                    p.add_run(f"  [Confidence: {item.get('confidence', 'medium')}]").italic = True
                    refs = ", ".join(item.get("evidence_ids", []))
                    if refs:
                        p.add_run(f" — {refs}").italic = True
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph("Not available.")

    # 8. Conflicting Evidence
    _heading("Conflicting Evidence")
    conflicts = report.get("conflicts", [])
    if conflicts:
        for c in conflicts:
            if isinstance(c, dict):
                doc.add_paragraph(
                    f"Type: {c.get('conflict_type', '?')} — {c.get('description', '')}"
                )
                doc.add_paragraph(f"Source A: {c.get('source_a_ref', '—')}", style="List Bullet")
                doc.add_paragraph(f"Source B: {c.get('source_b_ref', '—')}", style="List Bullet")
    else:
        doc.add_paragraph("No conflicting evidence detected.")

    # 9. Missing Data / Assumptions
    _heading("Missing Data / Assumptions")
    missing = report.get("missing_data", [])
    if missing:
        for item in missing:
            doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph("No missing data.")

    what_checked = report.get("what_was_checked", [])
    if what_checked:
        doc.add_heading("What Was Checked", level=1)
        for item in what_checked:
            doc.add_paragraph(str(item), style="List Bullet")

    required = report.get("required_data", [])
    if required:
        doc.add_heading("Required Data / Next Steps", level=1)
        for item in required:
            doc.add_paragraph(str(item), style="List Bullet")

    # 10. Sources
    _heading("Sources")
    sources = report.get("sources", [])
    if sources:
        src_table = doc.add_table(rows=1, cols=6)
        src_table.style = "Table Grid"
        for i, h in enumerate(["ID", "Type", "Title", "Reference", "Date", "Confidence"]):
            src_table.rows[0].cells[i].text = h
            src_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for src in sources:
            if isinstance(src, dict):
                row = src_table.add_row()
                row.cells[0].text = src.get("source_id", "—")
                row.cells[1].text = src.get("source_type", "—")
                row.cells[2].text = src.get("title", "—")
                row.cells[3].text = src.get("reference", "—")
                row.cells[4].text = src.get("date") or "—"
                row.cells[5].text = src.get("confidence", "—")
    else:
        doc.add_paragraph("No sources cited.")

    # 11. Quality Gate Status
    _heading("Quality Gate Status")
    doc.add_paragraph(f"Status: {qg_status}")

    doc.add_paragraph()
    footer = doc.add_paragraph(f"Generated by Decision Center — {generated_at}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
